
import asyncio
import json
from typing import List
from docx import Document
from discom.constants import JobStatus
from discom.queries import get_all_chunks, update_job, update_chunk
from db import get_consumer, get_db_pool, get_s3, init_db
from discom.constants import ChunkRecord
import io
import json
from settings import s3_settings
from asyncpg import Pool

def download_from_s3(s3_client, bucket: str, key: str) -> io.BytesIO:
    buffer = io.BytesIO()
    s3_client.download_fileobj(bucket, key, buffer)
    buffer.seek(0)
    return buffer


async def reassemble_document(chunks, s3_key, s3_client, bucket: str, pool: Pool, document_id: str):
    buffer = await asyncio.to_thread(download_from_s3, s3_client, bucket, s3_key)
    doc = Document(buffer)
    paragraphs: List[ChunkRecord] = [chunk for chunk in chunks if chunk.address["type"] == "paragraph"]
    tables: List[ChunkRecord] = [chunk for chunk in chunks if chunk.address["type"] == "table"]
    assembly_failures = 0

    for para in paragraphs:
        if para.status != JobStatus.DONE.value:
            continue
        try:
            doc.paragraphs[para.address["paragraph_index"]].text = para.result
        except Exception as e:
            assembly_failures += 1
            async with pool.acquire() as connection:
                await update_chunk(connection, JobStatus.ASSEMBLY_FAILED.value, para.id, str(e))
            

    for table in tables:
        if table.status != JobStatus.DONE.value:
            continue
        try:
            selected_table = doc.tables[table.address["table_index"]]
            cell_ids = table.address["cell_ids"]
            translated_cells = json.loads(table.result)
            if len(translated_cells) != len(cell_ids):
                assembly_failures += 1
                async with pool.acquire() as connection:
                    await update_chunk(connection, JobStatus.ASSEMBLY_FAILED.value, table.id, "cell count mismatch")
                continue

            for (row_idx, col_idx), translated_text in zip(cell_ids, translated_cells):
                selected_table.rows[row_idx].cells[col_idx].text = translated_text
        except Exception as e:
            assembly_failures += 1
            async with pool.acquire() as connection:
                await update_chunk(connection, JobStatus.ASSEMBLY_FAILED.value, table.id, str(e))

    return doc, assembly_failures

async def assemble_requests():
    consumer = get_consumer()
    pool = get_db_pool()
    s3 = get_s3()
    try:
        async for msg in consumer: 
            request = json.loads(msg.value.decode('utf-8'))
            document_id = request["document_id"]
            async with pool.acquire() as connection:
                chunks: List[ChunkRecord] = await get_all_chunks(connection, document_id)
            translated_doc, assembly_failures = await reassemble_document(chunks, document_id, s3, s3_settings.S3_BUCKET, pool, document_id)
            output_buffer = io.BytesIO()
            translated_doc.save(output_buffer)
            output_buffer.seek(0)
            s3_result_key = "translated/" + str(document_id)
            await asyncio.to_thread(s3.upload_fileobj, output_buffer, s3_settings.S3_BUCKET, s3_result_key)
            failed_chunks = [c for c in chunks if c.status in (JobStatus.FAILED.value, JobStatus.CANCELLED.value, JobStatus.ASSEMBLY_FAILED.value)]
            total_failures = len(failed_chunks) + assembly_failures
            final_status = JobStatus.DONE_WITH_ERRORS.value if total_failures else JobStatus.DONE.value
            error_msg = f"{total_failures} chunk(s) failed" if total_failures else None
            async with pool.acquire() as connection:
                await update_job(connection, final_status, document_id, s3_result_key, error_msg if total_failures else None)
            await consumer.commit()
    except Exception as e:
        print(f"Consumer loop crashed: {e}")
        async with pool.acquire() as connection:
            await update_job(connection, JobStatus.FAILED.value, document_id, None, str(e))
    finally:
        print("Stopping consumer...")
        await consumer.stop()
        await pool.close()


async def main():
    await init_db()
    await assemble_requests()

if __name__ == '__main__':
    asyncio.run(main())