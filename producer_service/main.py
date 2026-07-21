from contextlib import asynccontextmanager
from dataclasses import asdict
import io
import json
from uuid import uuid4
import asyncpg
from fastapi import Depends, FastAPI, Request, UploadFile, File
from discom.constants import RegisterRequest
from db.init_db import initialize_db
from discom.queries import add_job, register_user, get_user_id, bulk_insert_chunks, update_job
from db.connections import get_db_connection
from discom.constants import JobStatus
import redis.asyncio as redis
from middleware.rate_limiter import RateLimitMiddleWare
from middleware.api_auth import AuthMiddleWare
from settings import kafka_settings, s3_settings
from discom.constants import ChunkRecord
from docx import Document
import asyncio

@asynccontextmanager
async def lifespan(app: FastAPI):
    print("Starting up: Initializing database pools...")
    await initialize_db(app)
    yield
    print("Shutting down: Closing database pools...")
    await app.state.db_pool.close()
    await app.state.redis.aclose()
    await app.state.kafka_producer.stop()

app = FastAPI(lifespan=lifespan)
app.add_middleware(RateLimitMiddleWare)
app.add_middleware(AuthMiddleWare)

@app.get("/")
async def read_root():
    return {"Hello": "Aman Sheikh"}

def extract_chunks(document_id: str, docx_path) -> list[ChunkRecord]:
    doc = Document(docx_path)
    chunks: list[ChunkRecord] = []
    chunk_counter = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            chunk_counter += 1
            chunks.append(ChunkRecord(
                id=str(uuid4()),
                document_id=document_id,
                address={"type": "paragraph", "paragraph_index": i},
                source_text=text,
                status=JobStatus.PENDING.value,
                chunk_index=chunk_counter,
                result=""
            ))

    for t_idx, table in enumerate(doc.tables):
        non_empty = [
            (r, c, cell.text.strip())
            for r, row in enumerate(table.rows)
            for c, cell in enumerate(row.cells)
            if cell.text.strip()
        ]
        if non_empty:
            chunk_counter += 1
            chunks.append(ChunkRecord(
                id=str(uuid4()),
                document_id=document_id,
                address={"type": "table", "table_index": t_idx, "cell_ids": [(r, c) for r, c, _ in non_empty]},
                source_text=json.dumps([text for _, _, text in non_empty]),
                status=JobStatus.PENDING.value,
                chunk_index=chunk_counter,
                result=""
            ))
    return chunks

@app.post("/translate")
async def translate(request: Request, file: UploadFile = File(...), db: asyncpg.Connection = Depends(get_db_connection)):
    if file.content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return {
            "error": "Invalid file"
        }
    contents  = await file.read()
    document_id = str(uuid4())
    await asyncio.to_thread(app.s3.upload_fileobj, io.BytesIO(contents), s3_settings.S3_BUCKET, str(document_id))
    user_id = await get_user_id(db, request.headers.get("X-API-Key"))
    chunks = extract_chunks(document_id, io.BytesIO(contents))
    async with db.transaction():
        job_id = await add_job(db, user_id, document_id, file.filename, len(chunks), JobStatus.PENDING.value)
        await bulk_insert_chunks(db, chunks)

    try:
        async with app.state.kafka_producer.transaction():
            for chunk in chunks:
                payload = json.dumps({"chunk_id": chunk.id}).encode("utf-8")
                await app.state.kafka_producer.send_and_wait(
                    kafka_settings.KAFKA_INFERENCE_TOPIC,
                    payload
                )
    except Exception as e:
        async with db.transaction():
            await update_job(db, JobStatus.FAILED.value, document_id, None, f"Partial publish failure: {e}")
        return {"error": "Failed to queue document for translation", "detail": str(e)}
    return {
        "job_id": job_id,
        "status": "queued"
    }

@app.post("/register_user")
async def register(request: RegisterRequest, db: asyncpg.Connection = Depends(get_db_connection)):
    api_key = await register_user(db, request.user_name)
    return {
        "api_key": api_key
    }